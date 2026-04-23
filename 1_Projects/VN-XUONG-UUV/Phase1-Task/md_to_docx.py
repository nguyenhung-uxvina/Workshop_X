"""Convert Markdown files to DOCX using python-docx.
Handles: headings, tables, bold/italic, code blocks, bullet lists, horizontal rules.
"""
import re
import sys
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


def parse_inline(paragraph, text):
    """Parse inline markdown: **bold**, *italic*, `code`, ~~strikethrough~~."""
    # Split by inline patterns
    parts = re.split(r'(\*\*.*?\*\*|\*.*?\*|`[^`]+`|~~.*?~~)', text)
    for part in parts:
        if not part:
            continue
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
            run = paragraph.add_run(part[1:-1])
            run.italic = True
        elif part.startswith('`') and part.endswith('`'):
            run = paragraph.add_run(part[1:-1])
            run.font.name = 'Consolas'
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x80, 0x00, 0x00)
        elif part.startswith('~~') and part.endswith('~~'):
            run = paragraph.add_run(part[2:-2])
            run.font.strike = True
        else:
            paragraph.add_run(part)


def strip_inline_md(text):
    """Remove markdown inline formatting for plain text extraction."""
    text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
    text = re.sub(r'\*(.*?)\*', r'\1', text)
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'~~(.*?)~~', r'\1', text)
    return text


def convert_md_to_docx(md_path, docx_path):
    md_text = Path(md_path).read_text(encoding='utf-8')
    lines = md_text.split('\n')

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Strip YAML frontmatter
    i = 0
    if lines and lines[0].strip() == '---':
        i = 1
        while i < len(lines) and lines[i].strip() != '---':
            i += 1
        i += 1  # skip closing ---

    in_code_block = False
    code_lines = []
    in_table = False
    table_rows = []

    def flush_table():
        nonlocal table_rows, in_table
        if not table_rows:
            return
        # Parse columns from first row
        cols = [c.strip() for c in table_rows[0].strip('|').split('|')]
        num_cols = len(cols)

        # Skip separator row (row with ---)
        data_rows = []
        for row_str in table_rows:
            cells = [c.strip() for c in row_str.strip('|').split('|')]
            if cells and all(re.match(r'^[-:]+$', c) for c in cells if c):
                continue
            data_rows.append(cells)

        if not data_rows:
            table_rows = []
            in_table = False
            return

        table = doc.add_table(rows=len(data_rows), cols=num_cols)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        for r_idx, row_cells in enumerate(data_rows):
            for c_idx in range(min(num_cols, len(row_cells))):
                cell = table.cell(r_idx, c_idx)
                cell.text = ''
                p = cell.paragraphs[0]
                p.style = doc.styles['Normal']
                cell_text = strip_inline_md(row_cells[c_idx].strip())
                parse_inline(p, cell_text)
                p.paragraph_format.space_after = Pt(2)
                p.paragraph_format.space_before = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9)

            # Bold header row
            if r_idx == 0:
                for c_idx in range(min(num_cols, len(row_cells))):
                    for run in table.cell(r_idx, c_idx).paragraphs[0].runs:
                        run.bold = True

        doc.add_paragraph()  # spacing after table
        table_rows = []
        in_table = False

    while i < len(lines):
        line = lines[i]

        # Code block
        if line.strip().startswith('```'):
            if in_code_block:
                # End code block
                code_text = '\n'.join(code_lines)
                p = doc.add_paragraph()
                run = p.add_run(code_text)
                run.font.name = 'Consolas'
                run.font.size = Pt(9)
                p.paragraph_format.left_indent = Inches(0.3)
                code_lines = []
                in_code_block = False
            else:
                flush_table()
                in_code_block = True
                code_lines = []
            i += 1
            continue

        if in_code_block:
            code_lines.append(line)
            i += 1
            continue

        # Table row
        if '|' in line and line.strip().startswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            table_rows.append(line)
            i += 1
            continue
        else:
            if in_table:
                flush_table()

        stripped = line.strip()

        # Empty line
        if not stripped:
            i += 1
            continue

        # Horizontal rule
        if re.match(r'^[-*_]{3,}\s*$', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            # Add a thin line via border
            pPr = p._p.get_or_add_pPr()
            pBdr = pPr.makeelement(qn('w:pBdr'), {})
            bottom = pBdr.makeelement(qn('w:bottom'), {
                qn('w:val'): 'single',
                qn('w:sz'): '6',
                qn('w:space'): '1',
                qn('w:color'): 'auto',
            })
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # Headings
        heading_match = re.match(r'^(#{1,6})\s+(.*)', stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2)
            h = doc.add_heading(level=min(level, 4))
            parse_inline(h, text)
            i += 1
            continue

        # Blockquote
        if stripped.startswith('>'):
            text = re.sub(r'^>\s*', '', stripped)
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.4)
            parse_inline(p, text)
            for run in p.runs:
                run.italic = True
                run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            i += 1
            continue

        # Bullet list
        list_match = re.match(r'^(\s*)([-*+]|\d+\.)\s+(.*)', stripped)
        if list_match:
            text = list_match.group(3)
            bullet_char = list_match.group(2)
            if re.match(r'\d+\.', bullet_char):
                p = doc.add_paragraph(style='List Number')
            else:
                p = doc.add_paragraph(style='List Bullet')
            parse_inline(p, text)
            i += 1
            continue

        # Checkbox
        if stripped.startswith('- ['):
            checked = stripped.startswith('- [x]') or stripped.startswith('- [X]')
            text = re.sub(r'^- \[.\]\s*', '', stripped)
            marker = '☑' if checked else '☐'
            p = doc.add_paragraph()
            run = p.add_run(f'{marker} ')
            run.font.size = Pt(11)
            parse_inline(p, text)
            i += 1
            continue

        # Normal paragraph
        p = doc.add_paragraph()
        parse_inline(p, stripped)
        i += 1

    # Flush remaining table
    if in_table:
        flush_table()

    doc.save(docx_path)
    return docx_path


if __name__ == '__main__':
    base = Path(r'e:\Workshop_X\1_Projects\VN-XUONG-UUV\Phase1-Task')
    files = [
        'ICD_UUV_Interface_Template_v1.6_VN_updated.md',
        'SOP_Mother_Ship_Coordination_Template_v1.6_VN_updated.md',
        'SOP_Procurement_Acceptance_Template_v1.6_VN_updated.md',
        'SOP_UUV_Launch_Recovery_Template_v1.6_VN_updated.md',
    ]
    for f in files:
        md_path = base / f
        docx_path = base / f.replace('.md', '.docx')
        print(f'Converting: {f}')
        convert_md_to_docx(str(md_path), str(docx_path))
        print(f'  -> {docx_path.name}')
    print('Done.')
